"""Genera el reporte de actividades del mes 1 en formato .docx."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).resolve().parent.parent / "Reporte-actividades-mes-1.docx"

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
    return h

def add_paragraph_bold_lead(lead_bold, rest):
    p = doc.add_paragraph()
    run = p.add_run(lead_bold)
    run.bold = True
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_meta(label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)

# Título
title = doc.add_heading("Reporte de actividades — Mes 1", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title.runs:
    run.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)

add_meta("Periodo", "[insertar fechas del mes 1]")
add_meta("Programa", "Servicio social — CIIEMAD, Instituto Politécnico Nacional")
add_meta(
    "Líneas de trabajo",
    "Curso de bibliometría (artículo propio) · Observatorio de Arrecifes — México · "
    "Sección administrativa del Observatorio de Humedales",
)

doc.add_paragraph()

# Actividades realizadas
add_heading("Actividades realizadas", level=1)

actividades = [
    (
        "1. Inicio del curso de bibliometría",
        ", con el objetivo de elaborar un artículo de autoría propia sobre "
        "“Arquitectura verde asistida por inteligencia artificial” (primer planteamiento "
        "de tema). Durante el mes se revisaron metodologías de búsqueda y análisis de "
        "literatura científica, y se trabajó en la definición de la pregunta de "
        "investigación, los criterios de inclusión y exclusión de fuentes, y las "
        "primeras búsquedas en bases de datos académicas (Web of Science, Scopus, SciELO "
        "y Redalyc) sobre el uso de inteligencia artificial aplicada al diseño "
        "bioclimático, envolventes vegetales, techos verdes y certificaciones "
        "ambientales en edificación.",
    ),
    (
        "2. Arranque del proyecto “Observatorio de Arrecifes — México”",
        " como nueva iniciativa institucional del CIIEMAD-IPN, sumándose a los "
        "observatorios hermanos de humedales artificiales y techos verdes ya existentes. "
        "Se definió el alcance del observatorio como una plataforma web pública de "
        "monitoreo de 12 arrecifes coralinos mexicanos distribuidos en el Caribe "
        "(Puerto Morelos, Cozumel, Banco Chinchorro, Xcalak, Isla Contoy), Golfo de "
        "México (Sistema Arrecifal Veracruzano, Arrecife Alacranes) y Pacífico (Cabo "
        "Pulmo, Revillagigedo, Isla Isabel, Huatulco, Espíritu Santo).",
    ),
    (
        "3. Investigación documental sobre fuentes de datos",
        " que alimentarán la plataforma: agencias internacionales (NASA, NOAA, Agencia "
        "Espacial Europea, Servicio Geológico de Estados Unidos) e instituciones "
        "mexicanas (CONABIO, CONANP, INEGI, SEMARNAT). Se documentaron sus licencias de "
        "uso, frecuencia de actualización y formatos disponibles para asegurar que la "
        "información presentada en el observatorio respete las condiciones de cada "
        "proveedor y cite correctamente la fuente. Como referencias de diseño se "
        "identificaron dos plataformas internacionales: Allen Coral Atlas (mapas "
        "globales de arrecifes a alta resolución) y EJAtlas (atlas de conflictos "
        "socioambientales).",
    ),
    (
        "4. Construcción de la primera versión de la plataforma web",
        " del observatorio, incluyendo el sitio público con su página de inicio, "
        "identidad visual y estructura de navegación. Se diseñó una identidad visual "
        "océano-coral propia que distingue al observatorio de sus hermanos "
        "institucionales: paleta inspirada en los tonos del mar profundo y el coral "
        "vivo, tipografías combinadas para títulos y cuerpo de texto, y un sistema de "
        "tarjetas, botones y etiquetas reutilizables a lo largo del sitio.",
    ),
    (
        "5. Diseño de la portada del observatorio",
        " con una composición visual inspirada en los mapas batimétricos del Allen "
        "Coral Atlas, que evoca el fondo marino mediante capas de profundidad, "
        "isobatas, una retícula sutil que recuerda las celdas de monitoreo de hábitat "
        "y luces que simulan los reflejos de la superficie del agua. La portada "
        "incluye estadísticas animadas (número de arrecifes monitoreados, hectáreas "
        "bajo observación, etc.) que aparecen al cargar la página.",
    ),
    (
        "6. Inventario inicial de los 12 arrecifes",
        " con información completa para cada sitio: ubicación, superficie, "
        "profundidad, tipo de arrecife, estatus de protección (áreas naturales "
        "protegidas federales, sitios UNESCO, reservas de biósfera), porcentaje de "
        "cobertura de coral vivo, principales amenazas y riqueza de especies. "
        "Catálogo preliminar de 13 capas de información geográfica abierta que se "
        "integrarán al mapa del observatorio, todas con su atribución y licencia de "
        "uso correctamente documentadas.",
    ),
    (
        "7. Creación de la sección administrativa del Observatorio de Humedales "
        "Artificiales CDMX",
        ", que hasta ese momento sólo contaba con el sitio público. Se desarrolló el "
        "panel privado de gestión que permite al equipo del observatorio dar de alta, "
        "editar, archivar y publicar la información de los humedales y los hallazgos "
        "de monitoreo sin tener que tocar archivos de datos manualmente. Incluye "
        "pantalla de inicio de sesión, control de acceso por roles, un tablero general "
        "con resumen del estado del observatorio y formularios de captura para cada "
        "tipo de contenido. Esta sección administrativa sirvió además como base de "
        "referencia para el diseño del panel administrativo del Observatorio de "
        "Arrecifes que se construirá en el mes 2.",
    ),
]

for lead, rest in actividades:
    add_paragraph_bold_lead(lead, rest)

# Productos del mes
add_heading("Productos del mes", level=1)

productos_mes = [
    "Repositorio de código del Observatorio de Arrecifes creado y con primeras "
    "versiones funcionales corriendo en entorno local.",
    "Identidad visual y portada del sitio del Observatorio de Arrecifes terminadas.",
    "Inventario documentado de los 12 arrecifes.",
    "Catálogo de 13 fuentes de datos abiertas con sus respectivas licencias.",
    "Sección administrativa del Observatorio de Humedales lista para uso del equipo "
    "institucional.",
    "Bibliografía base inicial (~30 referencias) para el artículo de arquitectura "
    "verde con apoyo de IA.",
]
for item in productos_mes:
    doc.add_paragraph(item, style="List Bullet")

# Productos en proceso
add_heading("Productos en proceso (continúan en mes 2)", level=1)

productos_proceso = [
    "Conexión de la plataforma de Arrecifes con la base de datos institucional para "
    "cargar y administrar la información.",
    "Mapa interactivo con capas activables.",
    "Sección administrativa del Observatorio de Arrecifes (replicando el patrón del "
    "de humedales).",
    "Curaduría final del corpus bibliográfico para el artículo propio.",
]
for item in productos_proceso:
    doc.add_paragraph(item, style="List Bullet")

doc.save(OUTPUT)
print(f"Generado: {OUTPUT}")
